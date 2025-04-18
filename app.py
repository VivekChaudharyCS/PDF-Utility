from flask import Flask, request, jsonify, render_template, send_from_directory
import os
import time
import fitz
from io import BytesIO
from PIL import Image
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
PROCESSED_FOLDER = "processed_pdfs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/merge", methods=["POST"])
def merge_pdfs():
    if "files" not in request.files:
        return jsonify({"error": "No files part"}), 400

    files = request.files.getlist("files")
    if not files or len(files) < 2:
        return jsonify({"error": "Upload at least two PDF files"}), 400

    merger = PdfMerger()
    file_paths = []
    try:
        for file in files:
            file_path = os.path.join(UPLOAD_FOLDER, file.filename)
            file.save(file_path)
            file_paths.append(file_path)
            merger.append(file_path)

        output_filename = f"merged_{int(time.time())}.pdf"
        output_pdf = os.path.join(PROCESSED_FOLDER, output_filename)
        merger.write(output_pdf)
        merger.close()

        for file_path in file_paths:
            os.remove(file_path)

        return jsonify({"file_url": f"/download/{output_filename}"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/split", methods=["POST"])
def split_pdf():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    filename = f"{int(time.time())}_{file.filename}"
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(file_path)

    try:
        reader = PdfReader(file_path)
        total_pages = len(reader.pages)
        output_files = []

        for page_num in range(total_pages):
            writer = PdfWriter()
            writer.add_page(reader.pages[page_num])

            output_filename = f"split_page_{page_num + 1}_{int(time.time())}.pdf"
            output_path = os.path.join(PROCESSED_FOLDER, output_filename)

            with open(output_path, "wb") as output_pdf:
                writer.write(output_pdf)

            output_files.append(f"/download/{output_filename}")

        os.remove(file_path)

        return jsonify({"split_files": output_files})

    except Exception as e:
        return jsonify({"error": f"Failed to split PDF: {str(e)}"}), 500


@app.route("/compress-pdf", methods=["POST"])
def compress_pdf():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]

    if not file.filename.endswith(".pdf"):
        return jsonify({"error": "Invalid file format. Only PDFs are allowed."}), 400

    try:
        input_path = os.path.join(UPLOAD_FOLDER, file.filename)
        output_filename = f"compressed_{int(time.time())}.pdf"
        output_path = os.path.join(PROCESSED_FOLDER, output_filename)

        file.save(input_path)

        doc = fitz.open(input_path)

        for page in doc:
            images = page.get_images(full=True)

            for img in images:
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]

                img_pil = Image.open(BytesIO(image_bytes))
                img_pil = img_pil.convert("RGB")

                compressed_img_io = BytesIO()
                img_pil.save(compressed_img_io, format="JPEG", quality=50)

                compressed_img_bytes = compressed_img_io.getvalue()
                doc.update_stream(xref, compressed_img_bytes)

        doc.save(output_path, garbage=4, deflate=True)
        doc.close()

        os.remove(input_path)
        return jsonify({"file_url": f"/download/{output_filename}"})

    except Exception as e:
        print(e)
        return jsonify({"error": str(e)}), 500


@app.route("/extract_text", methods=["POST"])
def extract_text():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    if not file.filename.endswith(".pdf"):
        return jsonify({"error": "Invalid file format. Only PDFs are allowed."}), 400

    input_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(input_path)

    extracted_text = ""
    try:
        doc = fitz.open(input_path)
        for page in doc:
            extracted_text += page.get_text("text") + "\n\n"

        doc.close()
        os.remove(input_path)

        return jsonify({"extracted_text": extracted_text.strip()})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/reorder-pages", methods=["POST"])
def reorder_pages():
    if "file" not in request.files or "order" not in request.form:
        return jsonify({"error": "Missing file or order parameter"}), 400

    file = request.files["file"]
    order = request.form.get("order")

    if not file.filename.endswith(".pdf"):
        return jsonify({"error": "Invalid file format. Only PDFs are allowed."}), 400

    try:
        order_list = list(map(int, order.split(",")))

        input_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(input_path)

        reader = PdfReader(input_path)
        writer = PdfWriter()

        if any(page_num < 1 or page_num > len(reader.pages) for page_num in order_list):
            return jsonify({"error": "Invalid page numbers in order"}), 400

        for page_num in order_list:
            writer.add_page(reader.pages[page_num - 1])

        output_filename = f"reordered_{int(time.time())}.pdf"
        output_path = os.path.join(PROCESSED_FOLDER, output_filename)

        with open(output_path, "wb") as output_pdf:
            writer.write(output_pdf)

        os.remove(input_path)

        return jsonify({"file_url": f"/download/{output_filename}"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/rotate-pages", methods=["POST"])
def rotate_pages():
    if "file" not in request.files or "angles" not in request.form:
        return jsonify({"error": "Missing required parameters"}), 400

    file = request.files["file"]
    angles = request.form.get("angles")

    if not file.filename.endswith(".pdf"):
        return jsonify({"error": "Invalid file format. Only PDFs are allowed."}), 400

    try:
        rotation_angles = list(map(int, angles.split(",")))
        input_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(input_path)

        reader = PdfReader(input_path)
        writer = PdfWriter()

        for i, page in enumerate(reader.pages):
            for angle in rotation_angles:
                page.rotate(angle)
            writer.add_page(page)

        output_filename = f"rotated_{int(time.time())}.pdf"
        output_path = os.path.join(PROCESSED_FOLDER, output_filename)

        with open(output_path, "wb") as output_pdf:
            writer.write(output_pdf)

        os.remove(input_path)
        return jsonify({"file_url": f"/download/{output_filename}"})

    except Exception as e:
        print(e)
        return jsonify({"error": str(e)}), 500


@app.route("/delete-pages", methods=["POST"])
def delete_pages():
    if "file" not in request.files or "pages" not in request.form:
        return jsonify({"error": "Missing required parameters"}), 400

    file = request.files["file"]
    pages_to_delete = request.form.get("pages")

    if not file.filename.endswith(".pdf"):
        return jsonify({"error": "Invalid file format. Only PDFs are allowed."}), 400

    try:
        pages_to_delete = set(map(int, pages_to_delete.split(",")))
        input_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(input_path)

        reader = PdfReader(input_path)
        writer = PdfWriter()

        for i, page in enumerate(reader.pages, start=1):
            if i not in pages_to_delete:
                writer.add_page(page)

        output_filename = f"deleted_{int(time.time())}.pdf"
        output_path = os.path.join(PROCESSED_FOLDER, output_filename)

        with open(output_path, "wb") as output_pdf:
            writer.write(output_pdf)

        os.remove(input_path)
        return jsonify({"file_url": f"/download/{output_filename}"})

    except Exception as e:
        print(e)
        return jsonify({"error": str(e)}), 500


@app.route("/password-protect", methods=["POST"])
def password_protect():
    if "file" not in request.files or "password" not in request.form:
        return jsonify({"error": "Missing required parameters"}), 400

    file = request.files["file"]
    password = request.form.get("password").strip()

    if not file.filename.endswith(".pdf"):
        return jsonify({"error": "Invalid file format. Only PDFs are allowed."}), 400

    if not password:
        return jsonify({"error": "Password cannot be empty."}), 400

    try:
        input_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(input_path)

        reader = PdfReader(input_path)
        writer = PdfWriter()

        for page in reader.pages:
            writer.add_page(page)

        writer.encrypt(password)

        output_filename = f"protected_{int(time.time())}.pdf"
        output_path = os.path.join(PROCESSED_FOLDER, output_filename)

        with open(output_path, "wb") as output_pdf:
            writer.write(output_pdf)

        os.remove(input_path)
        return jsonify({"file_url": f"/download/{output_filename}"})

    except Exception as e:
        print(e)
        return jsonify({"error": str(e)}), 500


@app.route("/remove-password", methods=["POST"])
def remove_password():
    if "file" not in request.files or "password" not in request.form:
        return jsonify({"error": "Missing required parameters"}), 400

    file = request.files["file"]
    password = request.form.get("password")

    if not file.filename.endswith(".pdf"):
        return jsonify({"error": "Invalid file format. Only PDFs are allowed."}), 400

    try:
        input_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(input_path)

        reader = PdfReader(input_path)
        if not reader.is_encrypted:
            os.remove(input_path)
            return jsonify({"error": "The file is not password-protected."}), 400

        if not reader.decrypt(password):
            os.remove(input_path)
            return jsonify({"error": "Incorrect password."}), 400

        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)

        output_filename = f"unlocked_{int(time.time())}.pdf"
        output_path = os.path.join(PROCESSED_FOLDER, output_filename)

        with open(output_path, "wb") as output_pdf:
            writer.write(output_pdf)

        os.remove(input_path)
        return jsonify({"file_url": f"/download/{output_filename}"})

    except Exception as e:
        print(e)
        return jsonify({"error": str(e)}), 500


@app.route("/convert-pdf-to-docx", methods=["POST"])
def convert_pdf_to_docx():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]

    if not file.filename.endswith(".pdf"):
        return jsonify({"error": "Invalid file format. Only PDFs are allowed."}), 400

    try:
        input_path = os.path.join(UPLOAD_FOLDER, file.filename)
        output_filename = f"converted_{int(time.time())}.docx"
        output_path = os.path.join(PROCESSED_FOLDER, output_filename)

        file.save(input_path)

        doc = fitz.open(input_path)
        docx_file = Document()

        for page in doc:
            text = page.get_text("text")
            docx_file.add_paragraph(text)

        doc.close()
        docx_file.save(output_path)

        os.remove(input_path)

        return jsonify({"file_url": f"/download/{output_filename}"})

    except Exception as e:
        print(e)
        return jsonify({"error": str(e)}), 500


@app.route("/convert-docx-to-pdf", methods=["POST"])
def convert_docx_to_pdf():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]

    if not file.filename.endswith(".docx"):
        return jsonify(
            {"error": "Invalid file format. Only DOCX files are allowed."}
        ), 400

    try:
        input_path = os.path.join(UPLOAD_FOLDER, file.filename)
        output_filename = f"converted_{int(time.time())}.pdf"
        output_path = os.path.join(PROCESSED_FOLDER, output_filename)

        file.save(input_path)

        doc = Document(input_path)
        pdf_canvas = canvas.Canvas(output_path, pagesize=letter)
        y_position = 750

        for para in doc.paragraphs:
            pdf_canvas.drawString(100, y_position, para.text)
            y_position -= 20

            if y_position < 50:
                pdf_canvas.showPage()
                y_position = 750

        pdf_canvas.save()
        os.remove(input_path)

        return jsonify({"file_url": f"/download/{output_filename}"})

    except Exception as e:
        print(e)
        return jsonify({"error": str(e)}), 500


@app.route("/download/<filename>")
def download_file(filename):
    file_path = os.path.join(PROCESSED_FOLDER, filename)
    if os.path.exists(file_path):
        return send_from_directory(PROCESSED_FOLDER, filename, as_attachment=True)
    else:
        return jsonify({"error": "File not found"}), 404


if __name__ == "__main__":
    app.run(debug=True)
